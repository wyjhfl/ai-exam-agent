import os
import logging
from pathlib import Path
from pypdf import PdfReader
from docx import Document

logger = logging.getLogger(__name__)

CHUNK_SIZE = 500
CHUNK_OVERLAP = 50


class DocumentProcessor:

    @staticmethod
    def process_file(file_path: str) -> str:
        path = Path(file_path)
        ext = path.suffix.lower()
        if ext == ".pdf":
            return DocumentProcessor._extract_pdf(file_path)
        elif ext in (".docx", ".doc"):
            return DocumentProcessor._extract_docx(file_path)
        elif ext in (".txt", ".md"):
            return DocumentProcessor._extract_text(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        text_parts = []
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise
        return "\n\n".join(text_parts)

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        text_parts = []
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise
        return "\n\n".join(text_parts)

    @staticmethod
    def _extract_text(file_path: str) -> str:
        encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
        for enc in encodings:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    @staticmethod
    def chunk_text(text: str) -> list[str]:
        chunks = []
        start = 0
        while start < len(text):
            end = start + CHUNK_SIZE
            chunk = text[start:end]
            if chunk.strip():
                chunks.append(chunk.strip())
            start = end - CHUNK_OVERLAP
        return chunks
